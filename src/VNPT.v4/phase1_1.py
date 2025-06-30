import os, subprocess
from glob import glob
import docx
import re
from copy import deepcopy
import copy
import json
import pandas as pd
from asyncio.log import logger
import ipaddress
from datetime import datetime
from docx.shared import Pt
import logging
import sqlite3
import argparse
import os.path
import sys
import time
from docx.table import _Cell, Table
import numpy as np
from docx.shared import Inches
from docx.oxml import OxmlElement
import random
current_script_dir = os.path.dirname(os.path.abspath(__file__))
project_root_dir = os.path.abspath(os.path.join(current_script_dir, '..', '..'))
utils_dir_path = os.path.join(project_root_dir, 'utils')
if utils_dir_path not in sys.path:
    sys.path.insert(0, utils_dir_path)
from module_utils import *

def check_vietnamese(sn):
    list_tmp=sn.splitlines()
    list_SN=[]
    for i in list_tmp:
        if len(re.findall(r"\b\S*[ăâáắấàằầảẳẩãẵẫạặậđêéếèềẻểẽễẹệíìỉĩịôơóốớòồờỏổởõỗỡọộợưúứùừủửũữụựýỳỷỹỵ]+\S*\b", i.lower()))==0 and i.strip():
            list_SN.append(i)
    return list_SN

def convert_doc_to_docx(main_dir):
    for filename in os.listdir(main_dir):
        if filename.endswith('.doc'):
            print("Converting from doc to docx BBBG: "+filename)
            filename=os.path.join(main_dir, filename)
            subprocess.call(['libreoffice','--headless','--convert-to','docx',"--outdir", main_dir,f'{filename}'])
            # subprocess.call(['unoconv', '-d', 'document', '--format=docx', f'{filename}'])
    return glob(os.path.join(main_dir,"*.docx"))

def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_cell_background(cell, fill, color=None, val=None):
    """
    @fill: Specifies the color to be used for the background
    @color: Specifies the color to be used for any foreground
    pattern specified with the val attribute
    @val: Specifies the pattern to be used to lay the pattern
    color over the background color.
    """
    from docx.oxml.shared import qn
    from docx.oxml.xmlchemy import OxmlElement

    cell_properties = cell._element.tcPr
    try:
        cell_shading = cell_properties.xpath('w:shd')[0]  # in case there's already shading
    except IndexError:
        cell_shading = OxmlElement('w:shd') # add new w:shd element to it
    if fill:
        cell_shading.set(qn('w:fill'), fill)  # set fill property, respecting namespace
    if color:
        pass # TODO
    if val:
        pass # TODO
    cell_properties.append(cell_shading)  # finally extend cell props with shading element

def set_table_font(table, font_size, font_name):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = font_name
                    run.font.size = font_size

def get_SN_table_index(doc):
    for idx, table in enumerate(doc.tables):
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if 'Part #' in headers or 'Part#' in headers:
            return idx
    return -1

def parse_BBBG(folder_hd):
    """Parse BBBG folder from doc to docx and get the NET name and station name
    Args:
        folder_hd (string): path to folder
    Returns:
        dict: a dict with format {'ten hop dong': 'data'}
    """
    print("Parsing BBBG folder")
    hd=folder_hd.split("/")[-1]
    list_net=glob(folder_hd+"/*/", recursive = True)
    dict_bbbg_file={}
    dict_bbbg_file.update([(hd,[])])
    for net in list_net:
        list_file=convert_doc_to_docx(net)
        net=net.split('/')[-2]
        for file in list_file:
            print(f'Parsing bbbg {os.path.basename(file)}')
            head, tail=os.path.split(file)
            tail=re.search("(.+).docx",tail).group(1)
            file_copy=copy.deepcopy(file)
            wordDoc=docx.Document(file_copy)
            SN_table_index=get_SN_table_index(wordDoc)
            if SN_table_index==-1:
                print(f'No table DANH MỤC HÀNG HÓA BÀN GIAO TẠI TRẠM in BBBG {tail}')
                continue
            for para in wordDoc.paragraphs:
                if "tại trạm" in para.text.lower():
                    name_tram=re.search("tại trạm (.*)",para.text, re.I).group(1)
                    name_tram=f'Trạm {name_tram}'
                    break
            dict_bbbg_file[hd].append({'tail':tail, 'serial':{'fpc':[], 'mic':[], 'pic':[], 'module':[],'lca':[], 'chassis':[]}, 'net':net, 'name_tram':name_tram})
            table_header={}
            for i, cell in enumerate(wordDoc.tables[SN_table_index].rows[0].cells):
                table_header[cell.text]=i
            table_header['Part #']= table_header['Part #'] if 'Part #' in table_header else table_header['Part#']
            table_header={k.lower(): v for k, v in table_header.items()}
            if not all(header in table_header for header in ['serial number','part #','đvt','mô tả hàng hóa']):
                logging.exception('No column Serial Number/Part #/ĐVT/Mô tả hàng hóa in table Danh mục hàng hóa bàn giao in bbbg {}, list header is {}'.format(tail, table_header))
                print('No column Serial Number/Part #/ĐVT/Mô tả hàng hóa in table Danh mục hàng hóa bàn giao in bbbg {}'.format(tail))
                raise Exception
            throughput=None
            for row in wordDoc.tables[SN_table_index].rows:
                exception_element=["N/A", "", "Serial Number"]
                module=['XFP','QSFP','QDD','SFP']
                #########Part # ->>>> Mô tả hàng hóa
                if 'Linecard' in row.cells[table_header['mô tả hàng hóa']].text:
                    throughput= re.search("Linecard (.*)",row.cells[table_header['mô tả hàng hóa']].text, re.I).group(1)
                elif re.match(r"Card \d+G loại \d+", row.cells[table_header['mô tả hàng hóa']].text, re.I):
                    match=int(re.search(r"Card \d+G loại (\d+)", row.cells[table_header['mô tả hàng hóa']].text, re.I).group(1))
                    if match==2:
                        throughput='400G'
                    elif match==3:
                        throughput='200G'
                if all(row.cells[table_header['serial number']].text!=i for i in exception_element): #serial number not N/A or header
                    if row.cells[table_header['part #']].text.startswith('MPC') or row.cells[table_header['part #']].text.startswith('MX2K-MPC'):
                        SN=row.cells[table_header['serial number']].text
                        list_SN=check_vietnamese(SN)
                        dict_bbbg_file[hd][len(dict_bbbg_file[hd])-1]['serial']['fpc'].append({'listSN':list_SN,'PartNumber':row.cells[table_header['part #']].text,'Throughput':throughput})
                    elif row.cells[table_header['part #']].text=='MX2000-LC-ADAPTER':
                        SN=row.cells[table_header['serial number']].text
                        list_SN=check_vietnamese(SN)
                        dict_bbbg_file[hd][len(dict_bbbg_file[hd])-1]['serial']['lca'].append({'listSN':list_SN,'PartNumber':row.cells[table_header['part #']].text,'Throughput':None})
                    elif row.cells[table_header['part #']].text.startswith('MIC'):
                        SN=row.cells[table_header['serial number']].text
                        list_SN=check_vietnamese(SN)
                        dict_bbbg_file[hd][len(dict_bbbg_file[hd])-1]['serial']['mic'].append({'listSN':list_SN,'PartNumber':row.cells[table_header['part #']].text,'Throughput':None})
                    elif row.cells[table_header['part #']].text.startswith('PIC'):
                        SN=row.cells[table_header['serial number']].text
                        list_SN=check_vietnamese(SN)
                        dict_bbbg_file[hd][len(dict_bbbg_file[hd])-1]['serial']['pic'].append({'listSN':list_SN,'PartNumber':row.cells[table_header['part #']].text,'Throughput':None})
                    elif row.cells[table_header['part #']].text.startswith('MX960') or row.cells[table_header['part #']].text.startswith('MX2020'):
                        SN=row.cells[table_header['serial number']].text
                        list_SN=check_vietnamese(SN)
                        dict_bbbg_file[hd][len(dict_bbbg_file[hd])-1]['serial']['chassis'].append({'listSN':list_SN,'PartNumber':row.cells[table_header['part #']].text,'Throughput':None})
                    elif any(i in row.cells[table_header['part #']].text for i in module):
                        SN = row.cells[table_header['serial number']].text
                        list_SN=check_vietnamese(SN)
                        dict_bbbg_file[hd][len(dict_bbbg_file[hd])-1]['serial']['module'].append({'listSN':list_SN,'PartNumber':row.cells[table_header['part #']].text,'Throughput':None})
    return dict_bbbg_file

def validate_hostname(hostname):
    try:
        hostname=re.search("(.*)\DRE\d.*", hostname).group(1)
        return hostname
    except:
        return hostname

def validate_ip_address(address):
    try:
        ip = ipaddress.ip_address(address)
        return True
    except ValueError:
        return False

def validate_bbbg(bbbg):
    try:
        bbbg = re.search("(.+).(doc|docx)$", bbbg).group(1)
        return bbbg
    except:
        return bbbg

def parse_mapping(ip_file, mapping_file, output_dir, mapping_sheet="Sheet1", ip_sheet="Sheet1"):
    print("Parsing ip and mapping file")
    dict_ip_error={}
    if 'csv' in os.path.splitext(ip_file)[1].lower():
        df=pd.read_csv(ip_file, usecols=['IP Loopback', 'Hostname'])
    else:
        df=pd.read_excel(open(ip_file, 'rb'), sheet_name=ip_sheet,usecols=['IP Loopback', 'Hostname'])

    if 'csv' in os.path.splitext(mapping_file)[1].lower():
        df_mapping=pd.read_csv(mapping_file, usecols=['Hostname', 'BBBG'])
    else:
        df_mapping=pd.read_excel(open(mapping_file, 'rb'), sheet_name=mapping_sheet,usecols=['Hostname', 'BBBG'])
    # Validate data input - Start
    df['Hostname'] = df['Hostname'].apply(validate_hostname)
    df_mapping['Hostname'] = df_mapping['Hostname'].apply(validate_hostname)
    df_mapping['BBBG'] = df_mapping['BBBG'].apply(validate_bbbg)
    # Validate data input - End
    # Drop dupplicate to avoid dupplicate hostname, can cause error when use function "to_dict"
    error_data = df[df.isnull().any(axis=1)].drop_duplicates(subset=['Hostname'])
    not_null_data=df.dropna(how='any').dropna(how='all', axis=1)
    # Verify IP
    not_null_data['verify'] = not_null_data['IP Loopback'].apply(validate_ip_address)
    # Filter error data (dict_ip_error) and validated data (dict_ip_not_null)
    error_not_null = not_null_data[not_null_data['verify'] == False].drop_duplicates(subset=['Hostname'])
    validated = not_null_data[not_null_data['verify'] == True].drop_duplicates(subset=['Hostname'])
    validated=validated.rename(columns={"IP Loopback": "IP"})
    # Write error data to json file - Start
    dict_ip_error = pd.Series(error_data['IP Loopback'].values, index=error_data['Hostname']).to_dict()
    dict_ip_error.update(pd.Series(error_not_null['IP Loopback'].values, index=error_not_null['Hostname']).to_dict())
    with open(os.path.join(output_dir,"IP incorrect.json"), "w") as outfile:
        json.dump(dict_ip_error, outfile)
    # Write error data to json file - Stop
    df_mapping.drop_duplicates(subset=['Hostname'], inplace=True)
    return [validated[['IP', 'Hostname']], df_mapping[['Hostname', 'BBBG']]]

def add_random_minute_and_second(obj):
    if pd.isnull(obj):
        return obj
    random_minute = random.randint(0, 59)
    random_second = random.randint(0, 59)
    return obj.replace(hour=0, minute=random_minute, second=random_second)

def save_sqlite(output_dir, db_name, dict_bbbg_file, ip_df, mapping_df):
    print("Save to database")
    database = os.path.join(output_dir,db_name)
    conn = sqlite3.connect(database)
    bbbg_table = pd.merge(strip_df(ip_df), strip_df(mapping_df),  how='inner', left_on=['Hostname'], right_on = ['Hostname'])
    bbbg_df=pd.DataFrame()
    checkSn_table=pd.DataFrame()
    for key in dict_bbbg_file.keys():
        if not dict_bbbg_file[key]:
            continue
        df=pd.json_normalize(dict_bbbg_file[key], max_level=2)
        df['ma_HD']=key
        check=pd.DataFrame()
        for i in ['fpc','pic','mic','module','lca', 'chassis']:
            check=pd.concat([check,df[['tail','serial.{}'.format(i)]].assign(Type=i).rename(columns = {'serial.{}'.format(i):'serial'})])
        # check=pd.concat([df[['tail','serial.fpc']].assign(Type='fpc').rename(columns = {'serial.fpc':'serial'}),
        #                  df[['tail','serial.module']].assign(Type='module').rename(columns = {'serial.module':'serial'})])
        check=check.explode('serial')
        check = check.dropna(subset=["serial"])
        check=pd.concat([check.drop(['serial'], axis=1), check['serial'].apply(pd.Series)], axis=1).explode('listSN')
        check=check.rename(columns = {'listSN':'SN','tail':'BBBG'})
        check['TestStatus']='Unchecked'
        check['InstallationStatus']=None
        check['PlannedSlot']=None
        check['RealSlot']=None
        time_current=time.time()
        check['SN_create_timestamp']=time_current
        check['SN_status_update_timestamp']=time_current
        checkSn_table = check
        checkSn_table['Hostname']=''
        bbbg_df = df[['tail','net','ma_HD','name_tram']]
        checkSn_table['ma_HD'] = key
    list_bbbg= checkSn_table['BBBG'].unique()
    for bbbg in list_bbbg:
        host = ','.join(mapping_df.loc[mapping_df['BBBG'] == bbbg]['Hostname'].tolist())
        checkSn_table.loc[checkSn_table['BBBG'] == bbbg, 'Hostname'] = host
    bbbg_table = pd.merge(strip_df(bbbg_table), strip_df(bbbg_df), how='inner', left_on=['BBBG'], right_on=['tail']).drop(columns=['BBBG'])
    if bbbg_table.empty or checkSn_table.empty:
        print("BBBG table or SN is empty, check your input")
        logging.exception("Merging BBBG table or SN table is empty")
        raise Exception
    cur = conn.cursor()
    listOfTables = cur.execute(''' SELECT count(name) FROM sqlite_master WHERE type='table' AND name='checkSN' ''')

    if listOfTables.fetchone()[0]==1 :
        checkSn_table_db=pd.read_sql_query("SELECT TestStatus, InstallationStatus, PlannedSlot, RealSlot, Hostname, BBBG, SN, ma_HD, SN_status_update_timestamp FROM 'checkSN'" , conn)
        checkSN_result=pd.merge(strip_df(checkSn_table), strip_df(checkSn_table_db),  how='left', left_on=['SN', 'BBBG', 'ma_HD'], right_on = ['SN','BBBG', 'ma_HD'])
        for index, row in checkSN_result.iterrows():
            checkSN_result.loc[index,'SN_status_update_timestamp']=row['SN_status_update_timestamp_x']
            if not pd.isna(row['Hostname_y']) and not row['Hostname_y']=='' and row['TestStatus_y'] in ['Installed','Checked with reboot','Checked without reboot','Checked']:
                checkSN_result.loc[index,'TestStatus']=row['TestStatus_y']
                checkSN_result.loc[index,'InstallationStatus']=row['InstallationStatus_y']
                checkSN_result.loc[index,'PlannedSlot']=row['PlannedSlot_y']
                checkSN_result.loc[index,'RealSlot']=row['RealSlot_y']
                checkSN_result.loc[index,'Hostname']=row['Hostname_y']
                if row['TestStatus_y']!='Unchecked':
                    checkSN_result.loc[index,'SN_status_update_timestamp']=row['SN_status_update_timestamp_y']
            else:
                checkSN_result.loc[index,'TestStatus']=row['TestStatus_x']
                checkSN_result.loc[index,'InstallationStatus']=row['InstallationStatus_x']
                checkSN_result.loc[index,'PlannedSlot']=row['PlannedSlot_x']
                checkSN_result.loc[index,'RealSlot']=row['RealSlot_x']
                checkSN_result.loc[index,'Hostname']=row['Hostname_x']
        checkSN_result=checkSN_result.drop(columns=['TestStatus_x','InstallationStatus_x','PlannedSlot_x','RealSlot_x','Hostname_x','TestStatus_y','InstallationStatus_y','PlannedSlot_y','RealSlot_y','Hostname_y','SN_status_update_timestamp_y','SN_status_update_timestamp_x'])
        cur.execute("DELETE FROM 'checkSN' WHERE ma_HD = '{}'".format(list(dict_bbbg_file.keys())[0]))
    else:
        checkSN_result=checkSn_table
    checkSN_result['StatusTestStatus']='Valid'
    checkSN_result.loc[(checkSN_result['SN_status_update_timestamp']<checkSN_result['SN_create_timestamp']) & (checkSN_result['TestStatus']!='Unchecked'), 'StatusTestStatus'] = "Expired"

    bbbg_exist = cur.execute(''' SELECT count(name) FROM sqlite_master WHERE type='table' AND name='BBBG' ''')
    if bbbg_exist.fetchone()[0] == 1:
        cur.execute("DELETE FROM 'BBBG' WHERE ma_HD = '{}'".format(list(dict_bbbg_file.keys())[0]))
    sign_table=bbbg_table[['tail', 'name_tram', 'ma_HD', 'net']].rename({'tail': 'BBBG'}, axis=1).drop_duplicates()
    sign_table[['Người ký INOC trang 1', 'Người ký Netx trang 1', 'Người ký SVT trang 1', 'Người ký INOC chi tiết', 'Người ký SVT chi tiết', 'Ngày kết thúc', 'Thời gian ký', 'Người ký Netx chi tiết', 'Người ký Netx trang 1 ngoại quan','Người ký SVT trang 1 ngoại quan', 'Người ký Netx chi tiết ngoại quan','Người ký SVT chi tiết ngoại quan','Thời gian ký ngoại quan']] = None
    sign_result=sign_table
    sign_exist=cur.execute(''' SELECT count(name) FROM sqlite_master WHERE type='table' AND name='sign_time' ''')
    if sign_exist.fetchone()[0] == 1:
        sign_table_db=pd.read_sql_query("SELECT * FROM 'sign_time' where ma_HD=(?)" , conn, params=(list(dict_bbbg_file.keys())[0],))
        if not sign_table_db.empty:
            sign_result=pd.merge(strip_df(sign_table), strip_df(sign_table_db),  how='left', on=['BBBG', 'net','ma_HD'])
            x_columns = [col for col in sign_result.columns if col.endswith('_x')]
            new_columns = {
                x_col[:-2]: sign_result[x_col[:-2] + '_y'].combine_first(sign_result[x_col])
                for x_col in x_columns if x_col[:-2] + '_y' in sign_result.columns
            }
            sign_result = sign_result.assign(**new_columns)
            sign_result = sign_result.drop([col for col in sign_result.columns if col.endswith('_x') or col.endswith('_y')], axis=1)
            cur.execute("DELETE FROM 'sign_time' WHERE ma_HD = '{}'".format(list(dict_bbbg_file.keys())[0]))
    cur.close()
    bbbg_table.to_sql("BBBG", con=conn, schema=None, if_exists='append', index=False, index_label=None, chunksize=None, dtype=None, method=None)
    checkSN_result.to_sql("checkSN", con=conn, schema=None, if_exists='append', index=False, index_label=None, chunksize=None, dtype=None, method=None)
    sign_result.to_sql("sign_time", con=conn, schema=None, if_exists='append', index=False, index_label=None, chunksize=None, dtype=None, method=None)

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)

def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)

def delete_column_in_table(table, columns):
    grid = table._tbl.find("w:tblGrid", table._tbl.nsmap)
    for cell in table.column_cells(columns):
        cell._tc.getparent().remove(cell._tc)
    col_elem = grid[columns]
    grid.remove(col_elem)
    return grid

def get_first_table_after_heading(doc, heading_text):
    found_heading = False
    for block in doc.element.body:
        if block.tag.endswith('p'):
            para = block.text
            if heading_text in para:
                found_heading = True
        elif block.tag.endswith('tbl') and found_heading:
            return doc.tables[[t._element for t in doc.tables].index(block)]
    return None

def generate_atp(template, output_dir, hd, db_name, hopdong_dir):
    print("Generating atp hardware template")
    if template.endswith('.doc'):
        subprocess.call(['libreoffice','--headless','--convert-to','docx',"--outdir", os.path.dirname(template), f'{template}'])
        template=template.replace('.doc','.docx')
    database=os.path.join(output_dir,db_name)
    conn = sqlite3.connect(database)
    hd_dir=os.path.join(output_dir,hd)
    atp_dir=os.path.join(hd_dir,"ATP Template")
    cell_border_style =  {
                            "top": {"sz": 8, "val": "single"},
                            "bottom": {"sz": 8, "val": "single"},
                            "start": {"sz": 8, "val": "single"},
                            "end": {"sz": 8, "val": "single"},
                            }
    bbbg=pd.read_sql_query("SELECT tail, net, Hostname FROM 'BBBG' where ma_HD=(?)" , conn,params=(hd,))
    unique_bbbg_hd=pd.read_sql_query("SELECT * FROM 'sign_time' where ma_HD=(?)" , conn,params=(hd,))
    unique_bbbg_hd['inoc'] = unique_bbbg_hd['net'].str.extract(r'(?i)net\s*(\d+)', expand=False).where(lambda x: pd.notnull(x), None)
    unique_bbbg_hd['region'] = unique_bbbg_hd['inoc'].map({'1': 'Bắc','2': 'Nam','3': 'Trung'}).where(pd.notnull, None)
    unique_bbbg_hd['Thời gian ký'] = pd.to_datetime(unique_bbbg_hd['Thời gian ký']).dt.strftime('%d/%m/%y')
    listSN=pd.read_sql_query("SELECT SN, BBBG, PartNumber, Throughput, Type FROM 'checkSN' where ma_HD=(?)" , conn, params=(hd,))
    listSN['fpc_type_variable'] = listSN.apply(lambda row:
        "serial_number_here_MX960" if re.match(r"MX960(.*)", row['PartNumber']) else
        "serial_number_here_MX2020" if re.match(r"MX2020(.*)", row['PartNumber']) else
        'serial_number_here' +
        ('_' + re.search("MPC(.*?)-", row['PartNumber']).group(1) if re.search("MPC(.*?)-", row['PartNumber'])
        else '_' + re.search("MX2K-MPC(.*)", row['PartNumber']).group(1) if re.search("MX2K-MPC(.*)", row['PartNumber']) else '') +
        ('_' + row['Throughput'] if pd.notna(row['Throughput']) else ''),
        axis=1
    )
    # listSN['fpc_type_variable'] = '<serial_number_here'+listSN['PartNumber'].apply(lambda x: '_'+re.search("MPC(.*?)-",x).group(1) if re.search("MPC(.*?)-",x) else '').astype(str)+listSN['Throughput'].apply(lambda x:'_'+x if x is not None else '').astype(str)+'>'
    for index, unique_bbbg in unique_bbbg_hd.iterrows():
        print("Generating ATP hardware for BBBG: {}".format(unique_bbbg['BBBG']))
        bbbg_file=docx.Document(os.path.join(hopdong_dir, unique_bbbg['net'], unique_bbbg["BBBG"]+'.docx'))
        list_host=bbbg.loc[(bbbg['tail'] == unique_bbbg['BBBG']) & (bbbg['net'] == unique_bbbg['net'])]['Hostname'].unique()
        unique_bbbg['host_name']=', '.join(listSN.loc[listSN['BBBG'] == unique_bbbg['BBBG']]['PartNumber'].unique())
        atp_file = copy.deepcopy(docx.Document(template))
        has_fpc = any(type_val == 'fpc' for type_val in listSN.loc[listSN['BBBG'] == unique_bbbg['BBBG'], 'Type'].dropna())
        has_chassis = not listSN.loc[(listSN['BBBG'] == unique_bbbg['BBBG'])&(listSN['Type']=='chassis')].empty
        try:
            set_cell_text(tables=atp_file.tables,list_keyword=['host_name','name_tram', 'inoc', 'region', 'Người ký INOC trang 1', 'Người ký Netx trang 1', 'Người ký SVT trang 1', 'Người ký INOC chi tiết', 'Người ký SVT chi tiết', 'Người ký Netx chi tiết', 'Thời gian ký'], new_data=unique_bbbg)
            for table in atp_file.tables:
                if table.cell(0,0).paragraphs[0].text == '1_output_here' or table.cell(0,0).paragraphs[0].text == '2_output_here':
                    table._element.getparent().remove(table._element)
            for table in atp_file.tables:
                update_column_result = False
                for row in table.rows:
                    for cell in row.cells:
                        if '<serial_number_here' in cell.text:
                            sn_var=re.findall(r'<(\w*serial_number_here\w*)>',cell.text)[0]
                            # sn=', '.join(listSN.loc[(listSN['BBBG'] == unique_bbbg['BBBG'])&(listSN['Type'] == 'fpc')&(listSN['fpc_type_variable'] == sn_var)]['SN'].to_list())
                            sn=', '.join(listSN.loc[(listSN['BBBG'] == unique_bbbg['BBBG'])&(listSN['Type'] == 'fpc')&(listSN['fpc_type_variable'].str.contains(sn_var, na=False))]['SN'].to_list())
                            if sn=='':
                                cell.text='Trạm không được trang bị.'
                            else:
                                cell.text=cell.text.replace(f'<{sn_var}>',sn)
                            cell.paragraphs[0].runs[0].font.size = Pt(12)
                            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                        if ('show system license' in cell.text and not has_fpc) or (any(t in cell.text for t in ['Kiểm tra thành phần card điều khiển', 'Kiểm tra thành phần chuyển mạch', 'Kiểm tra thành phần nguồn', 'Kiểm tra thành phần Fantray', 'Kiểm tra thành phần Craft interface', 'khởi động lại card điều khiển']) and not has_chassis):
                            update_column_result=True
                    if update_column_result:
                        cell = row.cells[3]
                        for paragraph in cell.paragraphs:
                            cell._element.remove(paragraph._element)
                        paragraph = cell.add_paragraph('Không thực hiện mục này do phân bổ thành phần phần cứng tại trạm không có')
                        paragraph.runs[0].font.size = Pt(12)
                        paragraph.runs[0].font.name = 'Times New Roman'
                        paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
            tmp=0
            for item in atp_file.paragraphs:
                if "Kết quả test" in item.text:
                    tmp+=1
                    if ('510-2024' in hd or '117-2025' in hd) and (tmp<6 or tmp==7) and not has_chassis:
                        new_table=atp_file.add_table(rows=1, cols=1)
                        set_cell_border( cell = new_table.rows[0].cells[0], **cell_border_style)
                        heading_row = new_table.rows[0].cells
                        heading_row[0].text = 'No chassis'
                        run = new_table.cell(0, 0).paragraphs[0].runs[0]
                        run.font.name='Times New Roman'
                        run.font.size=Pt(12)
                        move_table_after(new_table, item)
                    elif ((('510-2024' in hd or '117-2025' in hd) and tmp==8) or ('510-2024' not in hd and '117-2025' not in hd and tmp==2)) and not has_fpc:
                        item._element.getparent().remove(item._element)
                    else:
                        for host in list_host:
                            new_table=atp_file.add_table(rows=1, cols=1)
                            set_cell_border( cell = new_table.rows[0].cells[0], **cell_border_style)
                            heading_row = new_table.rows[0].cells
                            heading_row[0].text = 'Output-'+str(tmp)+'-'+host
                            run = new_table.cell(0, 0).paragraphs[0].runs[0]
                            run.font.name='Times New Roman'
                            run.font.size=Pt(12)
                            move_table_after(new_table, item)
                            p1=item._p
                            p2=atp_file.add_paragraph(host)
                            p2.style = 'List'
                            p1.addnext(p2._p)
                elif "<input_table>" in item.text:
                    SN_table_index=get_SN_table_index(bbbg_file)
                    if SN_table_index==-1:
                        print(f'No table DANH MỤC HÀNG HÓA BÀN GIAO TẠI TRẠM in BBBG {unique_bbbg["BBBG"]}')
                        continue
                    table_component=bbbg_file.tables[SN_table_index]
                    table_component.autofit = False
                    table_component.preferred_width = Inches(6.2)
                    table_component.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
                    co_column=-1
                    csht_column=-1
                    for i, cell in enumerate(table_component.rows[0].cells):
                        if cell.text=='CO':
                            co_column=i
                        elif cell.text=='Mã CSHT':
                            csht_column=i
                    for i in sorted([csht_column,co_column],key=int, reverse=True):
                        if i!=-1:
                            delete_column_in_table(table_component,i)
                    for row_index in range(len(table_component.rows)):
                        for cell_index in range(len(table_component.rows[row_index].cells)):
                            set_cell_border( cell = table_component.rows[row_index].cells[cell_index] , **cell_border_style)
                    set_table_font(table=table_component, font_name='Times New Roman', font_size=Pt(12))
                    tbl_component = deepcopy(table_component._tbl)
                    paragraph=item.insert_paragraph_before()
                    paragraph._p.addnext(tbl_component)
                    item.text=""
            pre_file_name ="ATP_"+unique_bbbg["BBBG"]+".docx"
            atp_file.save(os.path.join(atp_dir,pre_file_name))
        except Exception as ex:
            print("Exception generating atp bbbg {}::: {}".format(unique_bbbg["BBBG"], ex))
            logging.exception(ex)
            raise Exception()

def PARSE_ARGS():
    """Parse command-line args"""
    parser = argparse.ArgumentParser(description='\nDemo reading argument when running script')
    INIT_LOGGING_ARGS(parser)
#====================================================
    parser.add_argument(
                '-hd',
                '--hopdong',
                type=str,
                help='\n\t\tFolder hop dong')

    parser.add_argument(
                '-m',
                '--mapping',
                type=str,
                help='\n\t\tFile mapping')

    parser.add_argument(
                '-ms',
                '--mapping_sheet',
                type=str,
                default="Sheet1",
                help='\n\t\tMapping sheet')

    parser.add_argument(
                '-i',
                '--ip',
                type=str,
                help='\n\t\tPlanning IP')

    parser.add_argument(
                '-is',
                '--ip_sheet',
                default="Sheet1",
                type=str,
                help='\n\t\tPlanning IP sheet')

    # parser.add_argument(
    #             '-s',
    #             '--signning',
    #             type=str,
    #             help='\n\t\tsignning file')

    # parser.add_argument(
    #             '-ss',
    #             '--signning_sheet',
    #             default="Sheet1",
    #             type=str,
    #             help='\n\t\tsignning file sheet')

    parser.add_argument(
                '-o',
                '--output_dir',
                type=str,
                default="/opt/ATP_output_result",
                help='\n\t\tDirectory save file json')

    parser.add_argument(
                '-th',
                '--template_hw',
                type=str,
                help='\n\t\tTemplate card file')

    parser.add_argument(
                '-ta',
                '--template_appearance',
                type=str,
                help='\n\t\tTemplate appearance file')

    parser.add_argument(
                '-db',
                '--database_name',
                type=str,
                default="database.sqlite",
                help='\n\t\tDatabase full name')

    parser.add_argument(
                '-shell_output',
                '--shell_output',
                choices = [ 'YES' , 'NO' ] ,
                default="NO",
                help='\n\t\toutput debug log file stdout?')

    return parser.parse_args()

def read_bbbg_data():
    args = PARSE_ARGS ( )
    if not args.hopdong:
        sys.exit
    hd=args.hopdong.split("/")[-1].replace(" ","_")
    #   =========== LOG INITIATION SEQUENCE
    dt = datetime.now()
    seq = str(dt.strftime("%Y%m%d"))
    pre_file_name ="Phase1.1-"+hd
    log_file_name = ("{}.log".format(pre_file_name))
    from  distutils import util
    LOGGER_INIT( log_level = args.log_level ,
							log_file = log_file_name ,
							shell_output = util.strtobool(args.shell_output) ,
							print_log_init = True)
#   =========== MAIN OPERATION
    bbbg = parse_BBBG(args.hopdong)
    ip_df, mapping_df = parse_mapping(args.ip, args.mapping, args.output_dir, args.mapping_sheet, args.ip_sheet)
    save_sqlite(args.output_dir, args.database_name, bbbg, ip_df, mapping_df)
    generate_atp(args.template_hw, args.output_dir, args.hopdong.split("/")[-1], args.database_name,args.hopdong)
    if args.template_appearance:
            generating_atp_appearance(args.hopdong.split("/")[-1], args.output_dir, args.database_name, args.template_appearance, args.hopdong)
    print('Done')

def generating_atp_appearance(hopdong, output_dir, database_name, template, hopdong_dir):
    from docx.oxml.ns import qn
    MAKE_DIR(os.path.join(output_dir, hopdong,'ATP Appearance'))
    if template.endswith('.doc'):
        subprocess.call(['libreoffice','--headless','--convert-to','docx',"--outdir", os.path.dirname(template), f'{template}'])
        template=template.replace('.doc','.docx')
    conn = sqlite3.connect(os.path.join(output_dir, database_name))
    unique_bbbgs_hd=pd.read_sql_query("SELECT * FROM 'sign_time' where ma_HD=(?)" , conn,params=(hopdong,))
    unique_bbbgs_hd['inoc'] = unique_bbbgs_hd['net'].str.extract(r'(?i)net\s*(\d+)', expand=False).where(lambda x: pd.notnull(x), None)
    unique_bbbgs_hd['region'] = unique_bbbgs_hd['inoc'].map({'1': 'Bắc','2': 'Nam','3': 'Trung'}).where(pd.notnull, None)
    listSN=pd.read_sql_query("SELECT BBBG, PartNumber, Type FROM 'checkSN' where ma_HD=(?)" , conn, params=(hopdong,))
    desired_columns = ['STT', 'Part #', 'Mô tả hàng hóa', 'ĐVT', 'SL', 'Serial Number']
    normalized_desired = [c.replace(' ', '').lower() for c in desired_columns]
    new_desired_columns=['STT','Ký mã hiệu sản phẩm','Tên hàng hóa và đặc tính kỹ thuật','ĐVT','SL','Serial Number']
    for index, unique_bbbg in unique_bbbgs_hd.iterrows():
        print("Generating ATP appearance for BBBG: {}".format(unique_bbbg['BBBG']))
        try:
            atp_hardware_template_file=docx.Document(os.path.join(hopdong_dir, unique_bbbg['net'], unique_bbbg["BBBG"]+'.docx'))
            atp_appearance_template_file = copy.deepcopy(docx.Document(template))
            unique_bbbg['host_name']=', '.join(listSN.loc[listSN['BBBG'] == unique_bbbg['BBBG']]['PartNumber'].unique())
            unique_bbbg['hardware']=', '.join(['linecard' if x.lower() == 'fpc' else x.lower() for x in listSN.loc[listSN['BBBG'] == unique_bbbg['BBBG']]['Type'].unique()])
            set_cell_text(tables=atp_appearance_template_file.tables,list_keyword=['host_name','name_tram', 'inoc', 'Người ký Netx trang 1 ngoại quan','Người ký SVT trang 1 ngoại quan', 'Người ký Netx chi tiết ngoại quan','Người ký SVT chi tiết ngoại quan','Thời gian ký ngoại quan', 'hardware'], new_data=unique_bbbg)
            for item in atp_appearance_template_file.paragraphs:
                if "<input_table>" in item.text:
                    SN_table_index=get_SN_table_index(atp_hardware_template_file)
                    if SN_table_index==-1:
                        print(f'No table DANH MỤC HÀNG HÓA BÀN GIAO TẠI TRẠM in ATP Template BBBG {unique_bbbg["tail"]}')
                        break
                    table_component=atp_hardware_template_file.tables[SN_table_index]
                    table_component.autofit = False
                    table_component.preferred_width = Inches(6.2)
                    table_component.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
                    header_texts = [cell.text.strip() for cell in table_component.rows[0].cells]
                    normalized_headers = [h.replace(' ', '').lower() for h in header_texts]
                    if not all(item in normalized_headers for item in normalized_desired):
                        missing = [col for col in desired_columns if col.replace(' ', '').lower() not in normalized_headers]
                        print(f'Missing column(s): {missing} in table DANH MỤC HÀNG HÓA BÀN GIAO TẠI TRẠM in ATP Template BBBG {unique_bbbg["tail"]}')
                        break
                    for i in reversed(range(len(header_texts))):
                        normalized = header_texts[i].replace(' ', '').lower()
                        if normalized not in normalized_desired:
                            delete_column_in_table(table_component, i)
                            del header_texts[i]
                    for cell, new_text in zip(table_component.rows[0].cells, new_desired_columns):
                        p = cell.paragraphs[0]
                        if p.runs:
                            r0 = p.runs[0]
                            name, size, bold, italic, underline = r0.font.name or "Times New Roman", r0.font.size or Pt(12), r0.font.bold or False, r0.font.italic or False, r0.font.underline or False
                        else:
                            name, size, bold, italic, underline = "Times New Roman", Pt(12), None, None, None
                        p.clear()
                        r = p.add_run(new_text)
                        f = r.font
                        f.name, f.size, f.bold, f.italic, f.underline = name, size, bold, italic, underline
                        r._element.rPr.rFonts.set(qn('w:eastAsia'), name)

                    tbl_component = copy.deepcopy(table_component._tbl)
                    set_table_font(table=Table(tbl_component, table_component._parent), font_name=name, font_size=size)
                    paragraph=item.insert_paragraph_before()
                    paragraph._p.addnext(tbl_component)
                    item.text=""
            atp_appearance_template_file.save(os.path.join(output_dir, hopdong,'ATP Appearance',"ATP_"+unique_bbbg["BBBG"]+".docx"))
        except Exception as ex:
            print("Exception generating atp appearance bbbg {}::: {}".format(unique_bbbg["BBBG"], ex))
            logging.exception(ex)
            raise Exception()

if __name__ == '__main__':
    read_bbbg_data()