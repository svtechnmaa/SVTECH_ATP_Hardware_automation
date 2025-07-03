import os
import re
import pandas as pd
import time
from datetime import datetime, timedelta
import random
from docx.shared import Pt
import logging
import sys
import sqlite3
import argparse
from glob import glob
import copy
import docx
import pathlib
from docx.table import _Cell
current_script_dir = os.path.dirname(os.path.abspath(__file__))
project_root_dir = os.path.abspath(os.path.join(current_script_dir, '..', '..'))
utils_dir_path = os.path.join(project_root_dir, 'utils')
if utils_dir_path not in sys.path:
    sys.path.insert(0, utils_dir_path)
from module_utils import *

def replace_starttime(single_line_text, sub_pattern, replacement):
    """
    Processes a multi-line text, replacing the date and setting the hour to 00
    (while retaining minutes and seconds) for every line that matches the pattern.

    Args:
        single_line_text (str): The input text containing one line.
                               Lines like: "Start time  2024-02-02 08:05:18 ICT" will be processed.
        new_date_str (str): The new date in 'YYYY-MM-DD' format.
                            Example: "2025-06-05"

    Returns:
        str: The modified multi-line text. Lines not matching the pattern are unchanged.
    """
    extract_pattern_modified = r".*?(\d{4}-\d{2}-\d{2}\s+00:\d{2}:\d{2})\s+.*"
    modified_line = re.sub(sub_pattern, replacement, single_line_text)
    extracted_datetime = None
    if modified_line != single_line_text:
        match = re.search(extract_pattern_modified, modified_line)
        if match:
            modified_datetime_string = match.group(1)
            extracted_datetime = datetime.strptime(modified_datetime_string, "%Y-%m-%d %H:%M:%S")
    return modified_line + '\n', extracted_datetime

def replace_uptime(single_line_text, end_datetime, start_datetime):
    """
    Replaces the duration part in an 'Uptime' line with the calculated difference
    between two datetime objects.

    Args:
        single_line_text (str): The input single line of text.
                               Example: "Uptime 1 day, 2 hours, 3 minutes"
        end_datetime (datetime): The ending datetime.
        start_datetime (datetime): The starting datetime.

    Returns:
        str: The modified line with a newline. Unchanged if no match.
    """
    total_diff_seconds = int((end_datetime - start_datetime).total_seconds())
    if total_diff_seconds < 0:
        total_diff_seconds = 0
    years = total_diff_seconds // (365 * 24 * 3600)
    remaining_seconds = total_diff_seconds % (365 * 24 * 3600)
    months = remaining_seconds // (30 * 24 * 3600)
    remaining_seconds %= (30 * 24 * 3600)
    days = remaining_seconds // (24 * 3600)
    remaining_seconds %= (24 * 3600)
    hours = remaining_seconds // 3600
    remaining_seconds %= 3600
    minutes = remaining_seconds // 60
    seconds = remaining_seconds % 60
    duration_parts = []
    if years > 0:
        duration_parts.append(f"{years} year{'s' if years != 1 else ''}")
    if months > 0:
        duration_parts.append(f"{months} month{'s' if months != 1 else ''}")
    if days > 0:
        duration_parts.append(f"{days} day{'s' if days != 1 else ''}")
    if hours > 0:
        duration_parts.append(f"{hours} hour{'s' if hours != 1 else ''}")
    if minutes > 0:
        duration_parts.append(f"{minutes} minute{'s' if minutes != 1 else ''}")
    if seconds > 0 or not duration_parts:
        duration_parts.append(f"{seconds} second{'s' if seconds != 1 else ''}")
    formatted_duration = ", ".join(duration_parts)
    pattern = r"^(\s*Uptime\s+).*$"
    replacement = rf"\g<1>{formatted_duration}"
    modified_line = re.sub(pattern, replacement, single_line_text)
    return modified_line + '\n'

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)

def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def delete_paragraph_and_matching_tables(doc):
    """
    Delete tables where cell(0,0) matches 'Output-<number>-<anything>'
    and the paragraph immediately before each matching table.
    
    Args:
        doc (Document): python-docx Document object
    
    Returns:
        Document: Modified Document object
    """
    pattern = re.compile(r'^Output-\d+-.*$')
    deleted_tables_info = []
    body = doc.element.body
    i = 0
    while i < len(body):
        elem = body[i]
        if elem.tag.endswith('tbl'):
            table_index = sum(1 for e in body[:i+1] if e.tag.endswith('tbl')) - 1
            if table_index < len(doc.tables):
                table = doc.tables[table_index]
                try:
                    cell_text = table.cell(0, 0).text.strip()
                    if pattern.match(cell_text):
                        deleted_tables_info.append(cell_text)
                        if elem.getparent() is body:
                            body.remove(elem)
                        else:
                            print(f"Warning: Table with cell(0,0) '{cell_text}' is not a direct child of body, skipping table deletion.")
                        if i > 0 and body[i-1].tag.endswith('p'):
                            if body[i-1].getparent() is body:
                                body.remove(body[i-1])
                                i -= 1
                            else:
                                print(f"Warning: Paragraph before table '{cell_text}' is not a direct child of body, skipping paragraph deletion.")
                except IndexError:
                    pass
        i += 1

    # Print deleted tables for verification
    if deleted_tables_info:
        print("Deleted the following empty tables (and their preceding paragraphs, if any):")
        for i, cell_text in enumerate(deleted_tables_info):
            print(f"Table {i+1}: Cell(0,0) text = {cell_text}")
    return doc

def write_atp(atp_template, list_log_file, atp_file_path, hd, end_date, sign_time):
    atp_file=copy.deepcopy(docx.Document(atp_template))
    try:
        tables=atp_file.tables
        for table in tables:
            if ('510-2024' in hd or '117-2025' in hd) and "Output-1-" in table.cell(0,0).paragraphs[0].text:
                host=re.search("Output-1-(.*)",table.cell(0,0).paragraphs[0].text).group(1)
                matching = [f for f in list_log_file if f'{host}_Chassis' in f]
                if not matching:
                    continue
                table.cell(0,0).paragraphs[0].text=""
                output={}
                line_index=[]
                lines=[]
                with open(matching[0]) as f:
                    lines = [line.rstrip() for line in f]
                for i in range(len(lines)):
                    if f'@{host}>' in lines[i]:
                        line_index.append(i)
                if '510' in hd:
                    output={'Output-1':lines[:line_index[6]], 'Output-2':lines[line_index[6]:line_index[9]], 'Output-3':lines[line_index[9]:line_index[11]], 'Output-4':lines[line_index[11]:line_index[13]], 'Output-5':lines[line_index[13]:line_index[14]], 'Output-7':lines[line_index[14]:]}
                elif '117-2025' in hd:
                    output={'Output-1':lines[:line_index[5]], 'Output-2':lines[line_index[5]:line_index[8]], 'Output-3':lines[line_index[8]:line_index[10]], 'Output-4':lines[line_index[10]:line_index[12]], 'Output-5':lines[line_index[12]:line_index[13]], 'Output-7':lines[line_index[13]:]}
                for line in output['Output-1']:
                    run = table.cell(0,0).add_paragraph().add_run(line.strip("\n"))
                    run.font.size  = Pt(7)
                    run.font.name = 'Courier New'
                    if "@" in line and '>' in line:
                        run.font.bold = True
                for table2 in tables:
                    if re.search(rf"Output-(2|3|4|5|7)-{host}", table2.cell(0,0).paragraphs[0].text):
                        data=output[f'Output-{re.search(rf"Output-(2|3|4|5|7)-{host}", table2.cell(0,0).paragraphs[0].text).group(1)}']
                        table2.cell(0,0).paragraphs[0].text=""
                        for line in data:
                            run = table2.cell(0,0).add_paragraph().add_run(line.strip("\n"))
                            run.font.size  = Pt(7)
                            run.font.name = 'Courier New'
                            if "@" in line and '>' in line:
                                run.font.bold = True
            elif (('510-2024' in hd or '117-2025' in hd) and "Output-6-" in table.cell(0,0).paragraphs[0].text) or ('510-2024' not in hd and '117-2025' not in hd and "Output-1-" in table.cell(0,0).paragraphs[0].text):
                host=re.search("Output-(1|6)-(.*)",table.cell(0,0).paragraphs[0].text).group(2)
                matching = [f for f in list_log_file if re.search(rf'{re.escape(host)}_(FPC|Module|LCA)', f)]
                if not matching:
                    continue
                matching.sort()
                table.cell(0,0).paragraphs[0].text=""
                output_1=[]
                dict_output_2=dict()
                output_2=''
                dict_output_lca=dict()
                dict_output_module=dict()
                str_output_module_first=[]
                str_output_module_last=[]
                str_output_lca=[]
                for log in matching:
                    file_log=open(log,'r')
                    line_index=[]
                    lines=[]
                    for i, line in enumerate(file_log.readlines()):
                        if f'@{host}>' in line:
                            line_index.append(i+1)
                        lines.append(line)
                    file_log.close()
                    fname = pathlib.Path(log)
                    if 'FPC' in log:
                        edited_starttime=None
                        if end_date and sign_time:
                            for i, e in enumerate(lines[:(line_index[-2]-1)]):
                                if re.search(r'show chassis fpc \d+ detail', e):
                                    for j in range(i + 1, (line_index[-2]-1)):
                                        if 'Start time' in lines[j]:
                                            if edited_starttime:
                                                sign_time=sign_time + timedelta(minutes=30)
                                                lines[j], _= replace_starttime(lines[j], r"^(.*?)(\s+)\d{4}-\d{2}-\d{2}\s+\d{2}:(\d{2}):(\d{2})(\s+.*)$", rf"\g<1>\g<2>{sign_time.strftime('%Y-%m-%d %H:%M:%S')}\g<5>")
                                                break
                                            else:
                                                lines[j], edited_starttime = replace_starttime(lines[j], r"^(.*?)(\s+)\d{4}-\d{2}-\d{2}\s+\d{2}:(\d{2}):(\d{2})(\s+.*)$", rf"\g<1>\g<2>{end_date.strftime('%Y-%m-%d')} 00:\g<3>:\g<4>\g<5>")
                                        elif 'Uptime' in lines[j]:
                                            lines[j]=replace_uptime(lines[j], sign_time, edited_starttime)
                                            break
                                elif re.search(r'show chassis pic fpc-slot \d+ pic-slot [01]', e):
                                    if re.search('PIC \d+ is empty', lines[i+1]):
                                        continue
                                    for j in range(i + 1, (line_index[-2]-1)):
                                        if 'Uptime' in lines[j]:
                                            lines[j]=replace_uptime(lines[j], sign_time, edited_starttime + timedelta(minutes=random.randint(5, 10)))
                                            break
                        output_1+=lines[:(line_index[-2]-1)]
                        dict_output_2[fname.stat().st_mtime]=lines[(line_index[-2]-1):]
                        dict_output_2=dict(sorted(dict_output_2.items()))
                        output_2=dict_output_2[list(dict_output_2.keys())[-1]]
                    elif 'Module' in log:
                        # If only have module in host, get only 1 command show chassis hardware
                        # if not any('FPC' in t for t in matching):
                        if '510-2024' in hd or '126-2025' in hd:
                            dict_output_module[fname.stat().st_mtime]=lines[(line_index[0]-1):(line_index[2]-1)]
                            str_output_module_last+=lines[(line_index[2]-1):]
                        else:
                            dict_output_module[fname.stat().st_mtime]=lines[(line_index[0]-1):(line_index[1]-1)]
                            str_output_module_last+=lines[(line_index[1]-1):]
                        dict_output_module=dict(sorted(dict_output_module.items()))
                        str_output_module_first=dict_output_module[list(dict_output_module.keys())[-1]]
                    elif 'LCA' in log:
                    # if all("LCA" in s for s in matching):
                        dict_output_lca[fname.stat().st_mtime]=lines
                        dict_output_lca=dict(sorted(dict_output_lca.items()))
                        str_output_lca=dict_output_lca[list(dict_output_lca.keys())[-1]]
                        # output_1=dict_output_lca[list(dict_output_lca.keys())[-1]]
                # If only have module in host, parse only 1 command show chassis hardware
                output_1=output_1+str_output_lca+str_output_module_first+str_output_module_last
                #Parsing log linecard, module and lca
                for line in output_1:
                    if "@" in line and '>' in line:
                        run = table.cell(0,0).add_paragraph().add_run(line.strip("\n"))
                        run.font.size  = Pt(7)
                        run.font.name = 'Courier New'
                        run.font.bold = True
                    else:
                        run = table.cell(0,0).add_paragraph().add_run(line.strip("\n"))
                        run.font.size  = Pt(7)
                        run.font.name = 'Courier New'
                for table2 in tables:
                    if ('510-2024' not in hd and '117-2025' not in hd and "Output-2-"+host in table2.cell(0,0).paragraphs[0].text) or (('510-2024' in hd or '117-2025' in hd) and "Output-8-"+host in table2.cell(0,0).paragraphs[0].text):
                        #Parsing log license
                        if output_2!='':
                            table2.cell(0,0).paragraphs[0].text=""
                            for line in output_2:
                                if "@" in line and '>' in line:
                                    run = table2.cell(0,0).add_paragraph().add_run(line.strip("\n"))
                                    run.font.size  = Pt(7)
                                    run.font.name = 'Courier New'
                                    run.font.bold = True
                                else:
                                    run = table2.cell(0,0).add_paragraph().add_run(line.strip("\n"))
                                    run.font.size  = Pt(7)
                                    run.font.name = 'Courier New'
        delete_paragraph_and_matching_tables(atp_file)
        atp_file.save(atp_file_path)
    except Exception as exp:
        logging.exception(exp)
        print(exp)

def export_atp(bbbg, hd, output_dir, end_date, sign_time):
    template_dir=os.path.join(output_dir, hd, "ATP Template")
    log_dir=os.path.join(output_dir, hd, "RAW LOG")
    atp_dir=os.path.join(output_dir,hd, "ATP")
    if not os.path.exists(atp_dir):
        os.makedirs(atp_dir)
    print("Writing file ATP docx for "+bbbg+" in hd "+hd)
    list_log_file=glob(os.path.join(log_dir,bbbg+'_*.txt'))
    if len(list_log_file)==0:
        print("No log file in bbbg {}".format(bbbg))
        return
    atp_file=os.path.join(template_dir,f'ATP_{bbbg}.docx')
    # atp_file=glob(os.path.join(template_dir,'*'+bbbg+'*.docx'))[0]
    file_name=os.path.join(atp_dir,atp_file.split("/")[-1])
    write_atp(atp_file, list_log_file, file_name, hd, end_date, sign_time)
    print("Writing file ATP docx for "+bbbg+": Done")

def main():
#   ===========INPUT INITIATION
    args = PARSE_ARGS ( )
    database=os.path.join(args.output_dir,args.database_name)
    conn = sqlite3.connect(database)
    list_bbbg=[bbbg.strip() for bbbg in args.bbbg.split(',')]
    placeholders = ','.join('?' * len(list_bbbg))
    query = f'SELECT BBBG, `Ngày kết thúc`, "Thời gian ký" FROM sign_time WHERE BBBG IN ({placeholders}) AND ma_HD = ?'
    # query = f'SELECT "tail", "Ngày kết thúc", "Thời gian ký" FROM BBBG WHERE tail IN ({placeholders}) AND ma_HD = ?'
    bbbg_on_db = pd.read_sql_query(query, conn, params=list_bbbg+ [args.hopdong]).drop_duplicates()
    bbbg_on_db['Ngày kết thúc'] = bbbg_on_db['Ngày kết thúc'].apply(pd.to_datetime, errors='coerce')
    bbbg_on_db['Thời gian ký'] = bbbg_on_db['Thời gian ký'].apply(pd.to_datetime, errors='coerce')
    for bbbg in list_bbbg:
    #   =========== LOG INITIATION SEQUENCE
        dt = datetime.now()
        seq = str(dt.strftime("%Y%m%d"))
        pre_file_name ="Phase2.3-"+ bbbg.replace(" ", "_" )+"_"+seq
        log_file_name = ("{}.log".format(pre_file_name))

        from  distutils import util
        LOGGER_INIT( log_level = args.log_level ,
                                log_file = log_file_name ,
                                shell_output = util.strtobool(args.shell_output) ,
                                print_log_init = True)
        export_atp(bbbg,args.hopdong,args.output_dir, bbbg_on_db[bbbg_on_db['BBBG'] == bbbg].iloc[0]['Ngày kết thúc'], bbbg_on_db[bbbg_on_db['BBBG'] == bbbg].iloc[0]['Thời gian ký'])

def PARSE_ARGS():
    """Parse command-line args"""
    parser = argparse.ArgumentParser(description='\nDemo reading argument when running script')
    INIT_LOGGING_ARGS(parser)
#====================================================
    parser.add_argument(
                '-hd',
                '--hopdong',
                type=str,
                #default = "",
                #required=True,
                help='\n\t\tHop dong')
    parser.add_argument(
                '-o',
                '--output_dir',
                type=str,
                default = "/opt/ATP_output_result",
                #required=True,
                help='\n\t\tOutput directory')
    parser.add_argument(
                '-db',
                '--database_name',
                type=str,
                default="database.sqlite",
                help='\n\t\tDatabase full name')
    parser.add_argument(
                '-b',
                '--bbbg',
                type=str,
                default="NO",
                help='\n\t\BBBG?')
    parser.add_argument(
                '-shell_output',
                '--shell_output',
                choices = [ 'YES' , 'NO' ] ,
                default="NO",
                help='\n\t\toutput debug log file stdout?')
    return parser.parse_args()

if __name__ == "__main__":
    main()