import argparse
import os
import sqlite3
import logging
import subprocess
import pandas as pd
import docx
import copy
from docx.shared import Inches
from phase1_1 import set_cell_text, get_SN_table_index, delete_column_in_table
import sys

def PARSE_ARGS():
    """Parse command-line args"""
    parser = argparse.ArgumentParser(description='\nDemo reading argument when running script')
    INIT_LOGGING_ARGS(parser)
#====================================================
    parser.add_argument(
                '-hd',
                '--hopdong',
                type=str,
                required=True,
                help='\n\t\tHop dong')
    parser.add_argument(
                '-template',
                '--template_file',
                type=str,
                required=True,
                help='\n\t\tTemplate file')
    parser.add_argument(
                '-o',
                '--output_dir',
                type=str,
                default="/opt/ATP_output_result",
                help='\n\t\tOutput directory')
    parser.add_argument(
                '-db',
                '--database_name',
                type=str,
                default="database.sqlite",
                help='\n\t\tDatabase full name')
    parser.add_argument(
                '-shell_output',
                '--shell_output',
                choices = [ 'YES' , 'NO' ] ,
                default="NO",
                help='\n\t\toutput debug log file stdout?')
    return parser.parse_args()

def create_dir(dir):
    if os.path.isdir(dir):
        print("Directory exists.")
    else:
        os.makedirs(dir)
        print("Directory not exists. Created")

def generating_atp_appearance(hopdong, output_dir, database_name, template):
    from docx.oxml.ns import qn
    create_dir(os.path.join(output_dir, hopdong,'ATP Appearance'))
    if template.endswith('.doc'):
        subprocess.call(['libreoffice','--headless','--convert-to','docx',"--outdir", os.path.dirname(template), f'{template}'])
        template=template.replace('.doc','.docx')
    conn = sqlite3.connect(os.path.join(output_dir, database_name))
    bbbg=pd.read_sql_query("SELECT tail, net, name_tram FROM 'BBBG' where ma_HD=(?)" , conn,params=(hopdong,)).drop_duplicates()
    bbbg['inoc'] = bbbg['net'].str.extract(r'(?i)net\s*(\d+)', expand=False).where(lambda x: pd.notnull(x), None)
    listSN=pd.read_sql_query("SELECT BBBG, PartNumber, Type FROM 'checkSN' where ma_HD=(?)" , conn, params=(hopdong,))
    desired_columns = ['STT', 'Part #', 'Mô tả hàng hóa', 'ĐVT', 'SL', 'Serial Number']
    new_desired_columns=['STT','Ký mã hiệu sản phẩm','Tên hàng hóa và đặc tính kỹ thuật','ĐVT','SL','Serial Number']
    for index, unique_bbbg in bbbg.iterrows():
        print("Generating ATP appearance for BBBG: {}".format(unique_bbbg['tail']))
        try:
            atp_hardware_template_file=docx.Document(os.path.join(output_dir, hopdong, 'ATP Template', "ATP_"+unique_bbbg["tail"]+".docx"))
            atp_appearance_template_file = copy.deepcopy(docx.Document(template))
            unique_bbbg['host_name']=', '.join(listSN.loc[listSN['BBBG'] == unique_bbbg['tail']]['PartNumber'].unique())
            unique_bbbg['hardware']=', '.join(['linecard' if x.lower() == 'fpc' else x.lower() for x in listSN.loc[listSN['BBBG'] == unique_bbbg['tail']]['Type'].unique()])
            set_cell_text(tables=atp_appearance_template_file.tables,list_keyword=['host_name','name_tram', 'inoc', 'Người ký Netx trang 1', 'Người ký SVT trang 1', 'hardware'], new_data=unique_bbbg)
            for item in atp_appearance_template_file.paragraphs:
                if "<input_table>" in item.text:
                    SN_table_index=get_SN_table_index(atp_hardware_template_file)
                    if SN_table_index==-1:
                        print(f'No table DANH MỤC HÀNG HÓA BÀN GIAO TẠI TRẠM in ATP Template BBBG {unique_bbbg["tail"]}')
                        break
                    table_component=atp_hardware_template_file.tables[SN_table_index]
                    table_component.autofit = False
                    table_component.preferred_width = Inches(6.2)
                    table_component.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
                    header_texts = [cell.text.strip() for cell in table_component.rows[0].cells]
                    if not all(item in header_texts for item in desired_columns):
                        print(f'Missing column {[col for col in desired_columns if col not in header_texts]} in table DANH MỤC HÀNG HÓA BÀN GIAO TẠI TRẠM in ATP Template BBBG {unique_bbbg["tail"]}')
                        break
                    for i in reversed(range(len(header_texts))):
                        if header_texts[i] not in desired_columns:
                            delete_column_in_table(table_component, i)
                            del header_texts[i]
                    for cell, new_text in zip(table_component.rows[0].cells, new_desired_columns):
                        p = cell.paragraphs[0]
                        if p.runs:
                            r0 = p.runs[0]
                            name, size = r0.font.name, r0.font.size
                            bold, italic, underline = r0.font.bold, r0.font.italic, r0.font.underline
                        else:
                            name, size, bold, italic, underline = None, None, None, None, None
                        p.clear()
                        r = p.add_run(new_text)
                        f = r.font
                        f.name, f.size = name, size
                        f.bold, f.italic, f.underline = bold, italic, underline
                        r._element.rPr.rFonts.set(qn('w:eastAsia'), name)

                    tbl_component = copy.deepcopy(table_component._tbl)
                    paragraph=item.insert_paragraph_before()
                    paragraph._p.addnext(tbl_component)
                    item.text=""
            atp_appearance_template_file.save(os.path.join(output_dir, hopdong,'ATP Appearance',"ATP_"+unique_bbbg["tail"]+".docx"))
        except Exception as ex:
            print("Exception generating atp appearance bbbg {}::: {}".format(unique_bbbg["tail"], ex))
            logging.exception(ex)
            raise Exception()

def main():
#   ===========INPUT INITIATION
    args = PARSE_ARGS ( )
    pre_file_name ="Phase2.4-"+args.hopdong
    log_file_name = ("{}.log".format(pre_file_name))

    from  distutils import util
    LOGGER_INIT( log_level = args.log_level ,
							log_file = log_file_name ,
							shell_output = util.strtobool(args.shell_output) ,
							print_log_init = True)
    generating_atp_appearance(hopdong=args.hopdong, output_dir=args.output_dir, database_name=args.database_name, template=args.template_file)

if __name__ == "__main__":
    main()