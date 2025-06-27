import pandas as pd
import os
import sqlite3
import random
from glob import glob
from docx import Document
from pathlib import Path
import sys
current_script_dir = os.path.dirname(os.path.abspath(__file__))
project_root_dir = os.path.abspath(os.path.join(current_script_dir, '..', '..'))
utils_dir_path = os.path.join(project_root_dir, 'utils')
if utils_dir_path not in sys.path:
    sys.path.insert(0, utils_dir_path)
from module_utils import *
from phase1_1 import set_cell_text

def PARSE_ARGS():
    """Parse command-line args"""
    parser = argparse.ArgumentParser(description='\nDemo reading argument when running script')
    INIT_LOGGING_ARGS(parser)
#====================================================

    parser.add_argument(
                '-hd',
                '--hopdong',
                type=str,
                default="HĐ 400")

    parser.add_argument(
                '-s',
                '--signning',
                type=str,
                help='\n\t\tsignning file')

    parser.add_argument(
                '-ss',
                '--signning_sheet',
                default="Sheet1",
                type=str,
                help='\n\t\tsignning file sheet')

    parser.add_argument(
                '-o',
                '--output_dir',
                type=str,
                default="/opt/ATP_output_result",
                help='\n\t\tDirectory save file json')

    parser.add_argument(
                '-db',
                '--database_name',
                type=str,
                default="database.sqlite",
                help='\n\t\tDatabase full name')

    parser.add_argument(
                '-shell_output',
                '--shell_output',
                choices = [ 'YES' , 'NO' ] ,
                default="NO",
                help='\n\t\toutput debug log file stdout?')

    return parser.parse_args()

def add_random_minute_and_second(obj):
    if pd.isnull(obj):
        return obj
    random_minute = random.randint(0, 59)
    random_second = random.randint(0, 59)
    return obj.replace(hour=0, minute=random_minute, second=random_second)

def process_signning(hopdong, signning, output_dir, database_name, signning_sheet='Sheet1', header_index=0):
    print('Parsing signing file')
    cols=['Tên trạm trên HS/BB', 'VNPT Net X', 'Người ký INOC trang 1', 'Người ký Netx trang 1', 'Người ký SVT trang 1', 'Người ký INOC chi tiết', 'Người ký SVT chi tiết', 'Ngày kết thúc', 'Thời gian ký', 'Người ký Netx chi tiết', 'Người ký Netx trang 1 ngoại quan','Người ký SVT trang 1 ngoại quan', 'Người ký Netx chi tiết ngoại quan','Người ký SVT chi tiết ngoại quan','Thời gian ký ngoại quan']
    if 'csv' in os.path.splitext(signning)[1].lower():
        df=pd.read_csv(signning, header=header_index)
    else:
        df=pd.read_excel(open(signning, 'rb'), sheet_name=signning_sheet, header=header_index)
    for col in cols:
        if col not in df.columns:
            df[col] = None
    df = df[cols]
    df = df.dropna(how='all')
    if pd.api.types.is_datetime64_any_dtype(df['Thời gian ký ngoại quan']):
        df['Thời gian ký ngoại quan'] = df['Thời gian ký ngoại quan'].apply(
            lambda x: x.strftime('%d/%m/%Y') if pd.notnull(x) else None
        )
    for col in ['Ngày kết thúc', 'Thời gian ký']:
        df[col] = pd.to_datetime(df[col], format='%m/%d/%Y', errors='coerce')
    df['Thời gian ký'] = df['Thời gian ký'].apply(add_random_minute_and_second)
    strip_df(df)
    database=os.path.join(output_dir,database_name)
    conn=sqlite3.connect(database)
    bbbg_table=pd.read_sql_query("SELECT tail, name_tram, ma_HD, net FROM 'BBBG' where ma_HD=(?)" , conn, params=(hopdong,)).drop_duplicates()
    sign_table = pd.merge(strip_df(bbbg_table), df, how='left', left_on=['name_tram', 'net'], right_on=['Tên trạm trên HS/BB', 'VNPT Net X'])
    sign_table = sign_table.drop(columns=['Tên trạm trên HS/BB', 'VNPT Net X'])
    sign_table = sign_table.rename(columns={"tail": "BBBG"})
    print('Saving signning data')
    cur = conn.cursor()
    sign_exist = cur.execute(''' SELECT count(name) FROM sqlite_master WHERE type='table' AND name='sign_time' ''')
    if sign_exist.fetchone()[0] == 1:
        cur.execute("DELETE FROM 'sign_time' WHERE ma_HD = '{}'".format(hopdong))
    cur.close()
    sign_table.to_sql("sign_time", con=conn, schema=None, if_exists='append', index=False, index_label=None, chunksize=None, dtype=None, method=None)
    print('Writing atp hardware template')
    list_file_template=glob(os.path.join(output_dir, hopdong, 'ATP Template', '*.docx'))
    sign_table['Thời gian ký'] = pd.to_datetime(sign_table['Thời gian ký']).dt.strftime('%d/%m/%y')
    for file in list_file_template:
        bbbg_data=sign_table.loc[sign_table['BBBG'] == Path(file).stem.replace('ATP_', '')]
        if not bbbg_data.empty:
            print(f'Writing atp hardware template file {Path(file).stem}')
            doc = Document(file)
            set_cell_text(tables=doc.tables,list_keyword=['Người ký INOC trang 1','Người ký Netx trang 1','Người ký SVT trang 1','Người ký INOC chi tiết','Người ký SVT chi tiết','Người ký Netx chi tiết', 'Thời gian ký'], new_data=bbbg_data.iloc[0])
            doc.save(file)
    list_file_appearance=glob(os.path.join(output_dir, hopdong, 'ATP Appearance', '*.docx'))
    for file in list_file_appearance:
        bbbg_data=sign_table.loc[sign_table['BBBG'] == Path(file).stem.replace('ATP_', '')]
        if not bbbg_data.empty:
            print(f'Writing atp appearance template file {Path(file).stem}')
            doc = Document(file)
            set_cell_text(tables=doc.tables,list_keyword=['Người ký Netx trang 1 ngoại quan','Người ký SVT trang 1 ngoại quan', 'Người ký Netx chi tiết ngoại quan','Người ký SVT chi tiết ngoại quan','Thời gian ký ngoại quan'], new_data=bbbg_data.iloc[0])
            doc.save(file)
    print('Done')

def read_signning():
    args = PARSE_ARGS ( )
    #   =========== LOG INITIATION SEQUENCE
    pre_file_name ="Phase1.3-"+args.hopdong
    log_file_name = ("{}.log".format(pre_file_name))

    from  distutils import util
    LOGGER_INIT( log_level = args.log_level ,
							log_file = log_file_name ,
							shell_output = util.strtobool(args.shell_output) ,
							print_log_init = True)

#   =========== MAIN OPERATION
    process_signning(hopdong=args.hopdong, signning=args.signning, output_dir=args.output_dir, database_name=args.database_name, signning_sheet=args.signning_sheet, header_index=2)

if __name__ == '__main__':
    read_signning()